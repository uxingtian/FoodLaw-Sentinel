from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class UnsupportedDocumentError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedDocumentError(f"仅支持 {allowed} 文件")
    if suffix in {".txt", ".md"}:
        return decode_text(content)
    if suffix == ".pdf":
        return extract_pdf_text(content)
    if suffix == ".docx":
        return extract_docx_text(content)
    raise UnsupportedDocumentError("不支持的文件类型")


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_cells: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text)
    return "\n".join(paragraphs + table_cells)


def clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t\f\v]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def chunk_text(text: str, target_size: int = 780, max_size: int = 900, min_size: int = 120) -> list[str]:
    cleaned = clean_text(text)
    if not cleaned:
        return []

    sentences = split_sentences(cleaned)
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        if not sentence:
            continue
        if len(sentence) > max_size:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(split_long_sentence(sentence, target_size))
            continue
        if current and len(current) + len(sentence) > max_size:
            chunks.append(current.strip())
            current = sentence
        else:
            current += sentence
        if len(current) >= target_size:
            chunks.append(current.strip())
            current = ""
    if current.strip():
        if chunks and len(current) < min_size:
            chunks[-1] = f"{chunks[-1]}\n{current.strip()}"
        else:
            chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk.strip()]


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[。！？；!?;])\s*|\n+", text)
    return [part.strip() for part in parts if part.strip()]


def split_long_sentence(sentence: str, size: int) -> list[str]:
    return [sentence[i : i + size].strip() for i in range(0, len(sentence), size) if sentence[i : i + size].strip()]
